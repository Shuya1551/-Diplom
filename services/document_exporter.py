"""
Модуль для экспорта новостей в документы .docx и .xlsx.
"""

import os
from datetime import datetime
from docx import Document
from openpyxl import Workbook

def export_to_docx(news_data, filepath):
    """
    news_data: список словарей или кортежей (текст, дата, план, id)
    filepath: полный путь к сохраняемому файлу .docx
    """
    doc = Document()
    doc.add_heading('Сгенерированные новости', level=1)
    doc.add_paragraph(f'Дата экспорта: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    for item in news_data:
        # item может быть кортежем (id, text, gen_date, approved, rating) + название плана
        # Упростим: ожидаем словарь
        if isinstance(item, dict):
            text = item.get('text', '')
            date = item.get('date', '')
            plan_title = item.get('plan_title', '')
        else:
            # кортеж: (id, text, gen_date, approved, rating, plan_title)
            text = item[1]
            date = str(item[2])
            plan_title = item[5] if len(item) > 5 else ''
        doc.add_heading(f'Мероприятие: {plan_title}', level=2)
        doc.add_paragraph(f'Дата генерации: {date}')
        doc.add_paragraph(text)
        doc.add_page_break()
    doc.save(filepath)

def export_to_xlsx(news_data, filepath):
    wb = Workbook()
    ws = wb.active
    ws.title = "Новости"
    ws.append(["ID", "Мероприятие", "Дата генерации", "Текст новости", "Утверждено", "Рейтинг"])
    for item in news_data:
        if isinstance(item, dict):
            ws.append([
                item.get('id', ''),
                item.get('plan_title', ''),
                item.get('date', ''),
                item.get('text', ''),
                "Да" if item.get('approved') else "Нет",
                item.get('rating', '')
            ])
        else:
            # кортеж (id, text, gen_date, approved, rating, plan_title)
            ws.append([
                item[0], item[5] if len(item)>5 else '', str(item[2]), item[1], 
                "Да" if item[3] else "Нет", item[4]
            ])
    wb.save(filepath)